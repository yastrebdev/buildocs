import os

from aiogram import html, Router, F
from aiogram.types import Message
from aiogram.fsm.context import FSMContext

import keyboards.inline_keyboards as kb
from bot import bot
from states.create_document import CreateDocument, CreateNote, document
from docx.shared import Cm

from utils.speech_decoding import speach_recognition

router = Router()


@router.message(CreateDocument.document_name)
async def create_name_document(message: Message, state: FSMContext):
    if not message.text:
        await message.answer('Жду название 🤖')
    else:
        await state.update_data(document_name=message.text)
        await message.answer(f'Вы создали документ: {html.bold(message.text)} 📄',
                             reply_markup=kb.action_with_document)

        data = await state.get_data()
        document.add_heading(data['document_name'], 0)
        await state.clear()


@router.message(CreateNote.photo)
async def send_note_photo(message: Message, state: FSMContext):
    if not message.photo:
        await message.answer('Это не похоже на фото...')
    else:
        photo = message.photo[-1]
        file_info = await bot.get_file(photo.file_id)
        file_path = file_info.file_path

        if not os.path.exists('downloads/photo'):
            os.makedirs('downloads/photo')

        await bot.download_file(file_path, f"downloads/photo/{photo.file_id}.jpg")
        document.add_picture(f"downloads/photo/{photo.file_id}.jpg", width=Cm(15))
        await message.answer('Вы добавили фотографию.', reply_markup=kb.create_note)

        await state.clear()


@router.message(CreateNote.voice)
async def send_note_voice(message: Message, state: FSMContext):
    if not message.voice:
        await message.answer('Похоже вы не знаете как отправить голосовое?')
    else:
        voice = message.voice
        file_id = voice.file_id

        voice_data = await bot.get_file(file_id)
        file_path = voice_data.file_path

        if not os.path.exists('downloads/voices'):
            os.makedirs('downloads/voices')
        await bot.download_file(file_path, f'downloads/voices/{file_id}.ogg')
        await message.answer('Подождите, сообщение декодируется...')

        text = speach_recognition(voice=f'downloads/voices/{file_id}.ogg')
        document.add_paragraph(text)

        await message.answer('Сообщение декодировано и добавлено', reply_markup=kb.create_note)

        await state.clear()


@router.message(CreateNote.text)
async def send_note_text(message: Message, state: FSMContext):
    if not message.text:
        await message.answer('Это не похоже на текст...')
    else:
        document.add_paragraph(message.text)
        await message.answer('Текст добавлен', reply_markup=kb.create_note)

        await state.clear()