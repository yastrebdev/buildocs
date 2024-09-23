from aiogram import html, Router, F
from aiogram.types import Message, FSInputFile, CallbackQuery
from aiogram.filters import CommandStart, Command
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext

from docx import Document
import app.keyboards as kb

document = Document()

router = Router()

class CreateDocument(StatesGroup):
    document_name = State()


@router.message(CommandStart())
async def start(message: Message):
    await message.answer(f"Привет, {html.bold(message.from_user.full_name)}!")


@router.message(Command('newdoc'))
async def create_document(message: Message, state: FSMContext):
    await state.set_state(CreateDocument.document_name)
    await message.answer(text='Введите название документа')


@router.message(CreateDocument.document_name)
async def create_name_document(message: Message, state: FSMContext):
    await state.update_data(document_name=message.text)
    await message.answer(f'Вы создали документ: {html.bold(message.text)}', reply_markup=kb.action_with_document)

    data = await state.get_data()
    document.add_heading(data['document_name'], 0)


@router.callback_query(F.data == 'create_note')
async def create_note(callback: CallbackQuery):
    await callback.message.answer('Отправьте фото, voice или текст', reply_markup=kb.create_note)

# @router.message(F.PHOTO)
# async def send_note_block(message: Message):
#     photo = message.photo[-1]
#     file_info = await bot.get_file(photo.file_id)
#     file_path = file_info.file_path
#     await bot.download_file(file_path, f"downloads/{photo.file_id}.jpg")


@router.callback_query(F.data == 'get_document')
async def create_note(callback: CallbackQuery):
    await callback.message.answer('Идет сборка документа...')

    document.save('demo.docx')
    file = FSInputFile('demo.docx')
    await callback.message.answer_document(file)
    await callback.message.answer('Выберите действие', reply_markup=kb.action_with_finished_document)


@router.callback_query(F.data == 'finished')
async def finished(callback: CallbackQuery):
    await callback.message.answer('Сессия завершена')


# @dp.message()
# async def get_voice(message: Message):
#     voice = message.voice
#     file_id = voice.file_id
#
#     voice_data = await bot.get_file(file_id)
#     file_path = voice_data.file_path
#
#     await bot.download_file(file_path, f'voices/{file_id}.ogg')
#
#     text = None
#     if os.path.exists(f'voices/{file_id}.ogg'):
#         text = speach_recognition(voice=f'voices/{file_id}.ogg')
#
#     await message.answer(text=text)